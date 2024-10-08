import argparse
import os
import shutil
import logging
from langchain.document_loaders.pdf import PyPDFDirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings.ollama import OllamaEmbeddings
from langchain_community.document_loaders import TextLoader
from langchain.schema.document import Document
from langchain_chroma import Chroma  
from docx import Document as DocxDocument



# Constants
CHROMA_PATH = "database"
DATA_PATH = "source_docs"

# Configure logging for production
logging.basicConfig(
    format="%(asctime)s - %(levelname)s - %(message)s",
    level=logging.INFO
)

def get_embedding_function():
    """Return the embedding function using Ollama embeddings."""
    try:
        embeddings = OllamaEmbeddings(model="nomic-embed-text")
        return embeddings
    except Exception as e:
        logging.error(f"Error initializing embeddings: {e}")
        raise

def main():
    """Main function to reset and update the Chroma database."""
    parser = argparse.ArgumentParser(description="Manage Chroma database.")
    parser.add_argument("--reset", action="store_true", help="Reset the database.")
    args = parser.parse_args()

    if args.reset:
        logging.info("Clearing the database as per --reset flag.")
        clear_database()

    try:
        documents = load_documents()
        chunks = split_documents(documents)
        add_to_chroma(chunks)
    except Exception as e:
        logging.error(f"An error occurred during processing: {e}")
        raise

def load_documents():
    """Load documents from the specified directory, handling PDFs, DOCX, and TXT."""
    if not os.path.exists(DATA_PATH):
        logging.error(f"Source directory '{DATA_PATH}' does not exist.")
        raise FileNotFoundError(f"{DATA_PATH} not found.")
    
    logging.info(f"Loading documents from {DATA_PATH}.")
    documents = []

        # Load PDF files using PyPDFDirectoryLoader
    pdf_doc = load_pdf(DATA_PATH)
    documents.extend(pdf_doc)

    # Loop through files in the directory
    for filename in os.listdir(DATA_PATH):
        file_path = os.path.join(DATA_PATH, filename)
        
        # if filename.endswith('.pdf'):
        #     documents.extend(load_pdf(file_path))
        if filename.endswith('.docx'):
            documents.extend(load_docx(file_path))  # Handle .docx files
        elif filename.endswith('.txt'):
            documents.extend(load_txt(file_path))  # Handle .txt files
        else:
            logging.warning(f"Unsupported file type: {filename}")

    if not documents:
        logging.warning("No documents were found in the source directory.")
    
    return documents

def load_pdf(file_path):
    """Load PDF using PyPDFDirectoryLoader."""
    pdf_loader = PyPDFDirectoryLoader(file_path)
    return pdf_loader.load()

def load_docx(file_path):
    """Load and split a .docx file into chunks."""
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]  # Ignore empty paragraphs

    formatted_documents = []
    current_chunk_index = 0  # Initialize the chunk index

    for page_num in range(0, len(paragraphs), 10):  # 10 paragraphs per page
        chunk = "\n\n".join(paragraphs[page_num:page_num + 10])
        if chunk.strip():  # Only add non-empty chunks
            chunk_id = f"{file_path}:{(page_num // 10) + 1}:{current_chunk_index}"
            formatted_documents.append(
                Document(metadata={
                    'source': file_path,
                    'page': (page_num // 10) + 1,
                    'ID': chunk_id  # Assign unique ID here
                }, page_content=chunk)
            )
            current_chunk_index += 1  # Increment for each chunk created

    return formatted_documents

def load_txt(file_path):
    """Load a .txt file using the TextLoader."""
    try:
        # Load the .txt file with UTF-8 encoding
        txt_loader = TextLoader(file_path, encoding='utf-8')
        return txt_loader.load()
    except Exception as e:
        logging.error(f"Error loading TXT file '{file_path}': {e}")
        return []


def split_documents(documents):
    """Split the documents into chunks for processing."""
    if not documents:
        logging.warning("No documents to split.")
        return []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=200,
        length_function=len
    )
    logging.info(f"Splitting {len(documents)} documents into chunks.")
    return text_splitter.split_documents(documents)

import uuid  # Import the uuid library for generating unique IDs

def add_to_chroma(chunks):
    """Add new document chunks to the Chroma database."""
    if not chunks:
        logging.warning("No chunks to add to the database.")
        return
    
    # Load or create a Chroma database
    db = Chroma(
        persist_directory=CHROMA_PATH, embedding_function=get_embedding_function()
    )

    chunks_with_ids = calculate_chunk_ids(chunks)
    existing_items = db.get(include=[])
    existing_ids = set(existing_items["ids"])
    
    logging.info(f"🔍 Existing documents in DB: {len(existing_ids)}")

    # Filter new chunks that are not already in the DB
    new_chunks = []
    new_chunk_ids = []

    for chunk in chunks_with_ids:
        chunk_id = chunk.metadata.get("id")
        if chunk_id not in existing_ids:
            if chunk_id is None:  # Check if the chunk ID is missing
                chunk_id = str(uuid.uuid4())  # Generate a new random ID
                chunk.metadata["id"] = chunk_id  # Assign the new ID to the chunk

            new_chunks.append(chunk)
            new_chunk_ids.append(chunk_id)

    # Log counts to check for mismatches
    logging.info(f"Number of new chunks: {len(new_chunks)}")
    logging.info(f"Number of new chunk IDs: {len(new_chunk_ids)}")

    if new_chunks:
        if len(new_chunks) == len(new_chunk_ids):  # Ensure they match
            logging.info(f"➕ Adding {len(new_chunks)} new documents to the database.")
            db.add_documents(new_chunks, ids=new_chunk_ids)
        else:
            logging.error("Mismatch in number of chunks and IDs.")
    else:
        logging.info("🆗 No new documents to add.")




def calculate_chunk_ids(chunks):
    """Generate unique IDs for each document chunk."""
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source")
        page = chunk.metadata.get("page")
        

        # Handle missing metadata
        if source is None or page is None:
            logging.warning("Chunk is missing source or page metadata.")
            continue  # Skip this chunk if essential metadata is missing
        
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk_id = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id
        chunk.metadata["id"] = chunk_id

    return chunks

def clear_database():
    """Clear the existing Chroma database."""
    if os.path.exists(CHROMA_PATH):
        logging.info(f"Removing existing database at {CHROMA_PATH}.")
        shutil.rmtree(CHROMA_PATH)
    else:
        logging.info("No database found to clear.")

if __name__ == "__main__":
    main()
